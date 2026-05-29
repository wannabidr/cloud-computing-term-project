from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PLAN_DOCX = ROOT / "reports/plan/[클라우드컴퓨팅] 텀 프로젝트 1. 제안서 (팀 번호).docx"
HANDOUT_DOCX = ROOT / "reports/plan/[클라우드컴퓨팅] Agent-as-a-Service 팀원 개발 핸드아웃.docx"
DIAGRAM = ROOT / ".codex_tmp/agentaas_architecture.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(85, 85, 85)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GRID = "D9E2EC"
FONT_LATIN = "Calibri"
FONT_KO = "Apple SD Gothic Neo"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = GRID, size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def apply_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def normalize_paragraph(paragraph, after: float = 8, line_spacing: float = 1.333, align=None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        apply_font(run, 11, BLACK)


def clear_cell(cell) -> None:
    cell.text = ""
    if cell.paragraphs:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)


def add_para(target, text: str = "", *, size: float = 11, bold: bool = False, color: RGBColor = BLACK,
             after: float = 8, before: float = 0, line_spacing: float = 1.333,
             align=None, style: str | None = None):
    p = target.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        apply_font(r, size, color, bold)
    return p


def add_heading(target, text: str, level: int = 1):
    p = target.add_paragraph()
    if level == 1:
        size, color, before, after = 16, BLUE, 18, 10
    elif level == 2:
        size, color, before, after = 13, BLUE, 12, 6
    else:
        size, color, before, after = 12, DARK_BLUE, 8, 4
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    apply_font(r, size, color, True)
    return p


def add_bullet(target, text: str, after: float = 4):
    p = target.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    apply_font(r, 11, BLACK)
    return p


def add_number(target, text: str, after: float = 4):
    p = target.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    apply_font(r, 11, BLACK)
    return p


def add_callout(target, title: str, body: str):
    table = target.add_table(rows=1, cols=1)
    set_table_width(table)
    set_table_borders(table, color="C9D6E2")
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    set_cell_shading(cell, LIGHT_GRAY)
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    apply_font(r, 11, DARK_BLUE, True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    apply_font(r2, 10.5, BLACK)
    add_para(target, "", after=4)


def add_simple_table(target, headers: Sequence[str], rows: Sequence[Sequence[str]],
                     widths: Sequence[int] | None = None):
    table = target.add_table(rows=1, cols=len(headers))
    set_table_width(table)
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        if widths:
            set_cell_width(cell, widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        clear_cell(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        apply_font(r, 10.5, DARK_BLUE, True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell)
            if widths:
                set_cell_width(cell, widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            clear_cell(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(text)
            apply_font(r, 10, BLACK)
    add_para(target, "", after=6)
    return table


def style_existing_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    for style_name in ("Normal", "Body Text"):
        if style_name not in styles:
            continue
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
        style.font.size = Pt(11)
        style.font.color.rgb = BLACK
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.333
    for table in doc.tables:
        set_table_width(table)
        set_table_borders(table, color="DADCE0")
        for row in table.rows:
            for cell in row.cells:
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    normalize_paragraph(p, after=6, line_spacing=1.2)


def add_cell_content(cell, blocks: Iterable[tuple[str, str | list[str] | tuple]]) -> None:
    clear_cell(cell)
    first = True
    for kind, value in blocks:
        if kind == "h1":
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(str(value))
            apply_font(r, 13, BLUE, True)
        elif kind == "h2":
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(str(value))
            apply_font(r, 12, DARK_BLUE, True)
        elif kind == "p":
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.25
            r = p.add_run(str(value))
            apply_font(r, 10.5, BLACK)
        elif kind == "bullet":
            p = cell.add_paragraph(style="List Bullet")
            first = False
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.125)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.18
            r = p.add_run(str(value))
            apply_font(r, 10.3, BLACK)
        elif kind == "num":
            p = cell.add_paragraph(style="List Number")
            first = False
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.125)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(str(value))
            apply_font(r, 10.3, BLACK)
        elif kind == "table":
            headers, rows, widths = value
            add_simple_table(cell, headers, rows, widths)
            first = False
        elif kind == "image":
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(value), width=Inches(6.1))
            p.paragraph_format.space_after = Pt(8)


def make_architecture_diagram(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    w, h = 1800, 1120
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("/System/Library/Fonts/AppleSDGothicNeo.ttc", 42)
        font = ImageFont.truetype("/System/Library/Fonts/AppleSDGothicNeo.ttc", 28)
        font_small = ImageFont.truetype("/System/Library/Fonts/AppleSDGothicNeo.ttc", 23)
    except Exception:
        font_title = ImageFont.load_default()
        font = ImageFont.load_default()
        font_small = ImageFont.load_default()

    def box(x1, y1, x2, y2, title, body, fill="#F4F6F9", outline="#2E74B5"):
        d.rounded_rectangle([x1, y1, x2, y2], radius=24, fill=fill, outline=outline, width=4)
        d.text((x1 + 24, y1 + 20), title, fill="#0B2545", font=font)
        for i, line in enumerate(body):
            d.text((x1 + 24, y1 + 68 + 34 * i), line, fill="#303030", font=font_small)

    def arrow(x1, y1, x2, y2, label=""):
        d.line([x1, y1, x2, y2], fill="#1F4D78", width=5)
        # Arrow head for mostly horizontal/vertical lines.
        if abs(x2 - x1) >= abs(y2 - y1):
            direction = 1 if x2 >= x1 else -1
            pts = [(x2, y2), (x2 - 24 * direction, y2 - 14), (x2 - 24 * direction, y2 + 14)]
        else:
            direction = 1 if y2 >= y1 else -1
            pts = [(x2, y2), (x2 - 14, y2 - 24 * direction), (x2 + 14, y2 - 24 * direction)]
        d.polygon(pts, fill="#1F4D78")
        if label:
            lx = (x1 + x2) // 2
            ly = (y1 + y2) // 2
            d.rounded_rectangle([lx - 145, ly - 28, lx + 145, ly + 28], radius=10, fill="white", outline="#D9E2EC")
            d.text((lx - 128, ly - 16), label, fill="#1F4D78", font=font_small)

    d.text((60, 42), "Multi-Tenant Agent-as-a-Service 구성도", fill="#0B2545", font=font_title)
    box(60, 145, 390, 300, "사용자", ["API Token으로 호출", "workspace_id 선택"], "#FFFFFF")
    box(515, 120, 900, 330, "AaaS API Gateway", ["인증/권한 검증", "Rate limit / Audit", "사용자-워크스페이스 매핑"], "#E8EEF5")
    box(1030, 120, 1460, 330, "OpenClaw Gateway", ["공용 agent 정의 실행", "sessionKey 생성", "agent target routing"], "#F4F6F9")
    box(1030, 430, 1460, 620, "공용 Agent Definition", ["프롬프트, 도구 정책", "모델/런타임 설정", "팀 공통 운영 자산"], "#FFFFFF")
    box(430, 730, 820, 935, "Tenant Workspace A", ["사용자 A 파일/상태", "컨테이너에만 mount"], "#FFFFFF")
    box(1000, 730, 1390, 935, "Tenant Workspace B", ["사용자 B 파일/상태", "사용자 간 파일 격리"], "#FFFFFF")
    box(80, 755, 330, 930, "OAuth Profiles", ["user A -> profile A", "user B -> profile B"], "#FFFDF5", "#A67C00")
    box(1510, 350, 1740, 690, "Sandbox", ["per-run/session", "CPU/Memory limit", "Docker network", "사용량 기록"], "#F9FBFD")
    box(1510, 760, 1740, 940, "Model API", ["OpenAI/Codex", "사용자별 OAuth"], "#FFFFFF", "#4F7CAC")

    arrow(390, 220, 515, 220)
    arrow(900, 220, 1030, 220)
    arrow(1245, 330, 1245, 430)
    arrow(1460, 520, 1510, 520)
    arrow(1510, 835, 1390, 835)
    arrow(820, 835, 1000, 835)
    arrow(330, 840, 430, 840)
    arrow(1625, 690, 1625, 760)

    d.text((65, 1025), "핵심: agent 시스템은 공용화하되 workspace, OAuth credential, sandbox runtime은 사용자 단위로 격리한다.", fill="#333333", font=font_small)
    img.save(path)


def build_plan_doc() -> None:
    make_architecture_diagram(DIAGRAM)
    doc = Document()

    title_table = doc.add_table(rows=2, cols=1)
    set_table_width(title_table)
    set_table_borders(title_table, color="FFFFFF", size="0")
    title_table.cell(0, 0).text = "클라우드컴퓨팅"
    title_table.cell(1, 0).text = "텀 프로젝트 1: 제안서"
    for idx, cell in enumerate((title_table.cell(0, 0), title_table.cell(1, 0))):
        set_cell_margins(cell, top=80, bottom=80, start=120, end=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            apply_font(r, 16 if idx == 0 else 20, DARK_BLUE if idx == 0 else BLACK, True)

    add_para(doc, "", after=8)
    meta_table = doc.add_table(rows=4, cols=2)
    set_table_width(meta_table)
    set_table_borders(meta_table)
    for row, label in enumerate(["대표학생 이름", "대표학생 학번", "대표학생 소속 학과/대학", "분반"]):
        meta_table.cell(row, 0).text = label
        meta_table.cell(row, 1).text = ""
        set_cell_shading(meta_table.cell(row, 0), LIGHT_GRAY)
    add_para(doc, "", after=6)

    add_callout(
        doc,
        "작성 방향",
        "본 계획서는 OpenClaw를 기반으로 공용 agent 시스템, 사용자별 workspace, 사용자별 OAuth 인증, 컨테이너 격리 실행을 결합한 Multi-Tenant Agent-as-a-Service 플랫폼을 제안한다.",
    )

    def add_question(text: str):
        table = doc.add_table(rows=1, cols=1)
        set_table_width(table)
        set_table_borders(table)
        set_cell_shading(table.cell(0, 0), LIGHT_GRAY)
        table.cell(0, 0).text = text
        add_para(doc, "", after=2)
        return table

    def add_answer():
        table = doc.add_table(rows=1, cols=1)
        set_table_width(table)
        set_table_borders(table)
        add_para(doc, "", after=8)
        return table

    add_question("[1] 개발 내용 소개\n텀 프로젝트 주제 및 텀 프로젝트에서 개발하고자 하는 시스템/서비스/SW에 대해 소개하세요.")
    add_answer()
    add_answer()
    add_question("[2] 필요성 및 활용방안\n개발하고자 하는 SW의 필요성에 대해 설명하고, 개발 결과물이 어떻게 사용될지 Use-Case를 설명하세요.")
    add_answer()
    add_question("[3] 선행기술 조사\n동일하거나 유사한 기술을 조사하고, 관련 제품/기술의 간략한 소개 및 출처를 작성하세요.")
    add_answer()
    add_question("[4] 제안하는 SW/시스템의 구성도\n최종 결과물의 구성 및 동작방식을 다이어그램과 함께 설명하세요.")
    add_answer()
    add_question("[5] 동작 환경\n최종 결과물이 어느 플랫폼에서 동작하는지 설명하세요.")
    add_answer()

    style_existing_doc(doc)

    # Preserve template fields, but fill known representative metadata from the previous team draft.
    meta = doc.tables[1]
    meta.cell(0, 1).text = "김성욱"
    meta.cell(1, 1).text = "202024191"
    meta.cell(2, 1).text = "정보의생명공학대학 정보컴퓨터공학부"
    meta.cell(3, 1).text = "059"

    for row in meta.rows:
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                normalize_paragraph(p, after=2, line_spacing=1.1)

    subject = "사용자별 Workspace와 OAuth 인증을 지원하는 OpenClaw 기반 Multi-Tenant Agent-as-a-Service 플랫폼"
    doc.tables[4].cell(0, 0).text = subject
    for p in doc.tables[4].cell(0, 0).paragraphs:
        normalize_paragraph(p, after=4, line_spacing=1.2, align=WD_ALIGN_PARAGRAPH.CENTER)
        for r in p.runs:
            apply_font(r, 12, DARK_BLUE, True)

    intro = [
        ("p", "본 프로젝트는 OpenClaw를 로컬 개인 실행 도구가 아니라 퍼블릭 클라우드에서 상시 운영되는 공용 Multi-Agent-as-a-Service 플랫폼으로 확장하는 것을 목표로 한다. 사용자는 무거운 멀티에이전트 런타임을 각자 노트북에 설치하지 않고, API Gateway를 통해 팀이 공유하는 agent 시스템을 호출한다."),
        ("p", "서비스의 핵심 아이디어는 agent 정의와 실행 플랫폼은 공용화하되, 실제 작업 상태는 사용자별 workspace로 분리하는 것이다. 동일한 코드 리뷰 agent를 호출하더라도 사용자 A는 A의 workspace에서, 사용자 B는 B의 workspace에서 작업하며, 각 실행은 Docker sandbox container 안에서 수행된다."),
        ("p", "또한 OpenClaw가 제공하는 OpenAI/Codex OAuth 기반 모델 접근 기능을 재사용하되, 서비스 계층에서 사용자별 OAuth profile을 선택하는 Tenant-aware OAuth Routing 기능을 추가한다. 이를 통해 공용 agent 시스템을 사용하면서도 모델 사용 계정, workspace, 실행 컨테이너, 사용량 기록은 사용자 단위로 분리한다."),
        ("h2", "주요 개발 목표"),
        ("bullet", "Agent API Gateway: 외부 사용자가 특정 agent를 REST/OpenAI-compatible API 형태로 호출할 수 있는 엔드포인트를 제공한다."),
        ("bullet", "Tenant Workspace Router: 사용자 토큰과 workspace_id를 검증하고 서버 내부의 안전한 workspace 경로로 매핑한다."),
        ("bullet", "OAuth Profile Router: 사용자별 OpenAI/Codex OAuth profile을 선택하여 모델 호출이 각 사용자의 인증 범위에서 수행되도록 한다."),
        ("bullet", "Container Runtime Manager: agent 실행 도구와 workspace 접근을 Docker sandbox 안에서 수행하고 CPU, memory, network, idle cleanup 정책을 적용한다."),
        ("bullet", "Usage/Audit Monitor: 요청 수, 실행 시간, 컨테이너 리소스 사용량, 사용자별 agent 사용 이력을 기록한다."),
    ]
    add_cell_content(doc.tables[5].cell(0, 0), intro)

    need = [
        ("h1", "필요성"),
        ("p", "OpenClaw와 같은 멀티에이전트 시스템은 여러 agent, tool runtime, sandbox container, browser/tool process가 동시에 동작할 수 있어 로컬 PC에서 실행할 때 메모리와 CPU 부담이 크다. 팀 단위로 동일한 agent 시스템을 사용해야 하는 경우에는 각 구성원이 같은 설정을 반복 구축해야 하고, 개인 노트북의 성능 차이와 환경 차이로 인해 실행 결과와 운영 안정성이 달라질 수 있다."),
        ("p", "또한 로컬 실행은 상시 가동이 어렵다. 노트북이 꺼지거나 네트워크가 끊기면 agent 시스템도 중단되며, 회사나 팀에서 공통으로 쓰는 자동화 agent를 중앙에서 관리하기 어렵다. 본 프로젝트는 이러한 문제를 해결하기 위해 agent 시스템을 클라우드에 상시 배포하고, 사용자는 가벼운 API client만으로 공용 agent를 호출하는 구조를 제안한다."),
        ("h1", "활용방안 및 Use-Case"),
        ("bullet", "팀 공용 코드 리뷰 agent: 모든 팀원이 동일한 코드 리뷰 정책을 가진 agent를 사용하되, 각자 다른 repository workspace에서 결과를 받는다."),
        ("bullet", "문서 요약/보고서 작성 agent: 회사 또는 수업 팀이 공용 agent 정의를 관리하고, 사용자는 자신의 자료 workspace를 지정하여 요약 또는 초안을 생성한다."),
        ("bullet", "신규 구성원 온보딩: 로컬에 복잡한 OpenClaw 환경을 설치하지 않고 API token과 workspace만 발급받아 곧바로 agent를 사용할 수 있다."),
        ("bullet", "상시 자동화: 클라우드 Gateway가 24시간 실행되므로 예약 작업, 장시간 분석, 팀 공유 자동화를 안정적으로 수행할 수 있다."),
        ("bullet", "사용량 기반 비용 분석: 사용자별 요청 수, 컨테이너 실행 시간, CPU/memory 사용량을 기록하여 내부 비용 정산 또는 quota 정책의 근거로 활용한다."),
        ("h1", "기대효과"),
        ("bullet", "로컬 컴퓨터 자원 부담 감소: 사용자는 브라우저 또는 CLI/API client만 사용하고, 무거운 agent 실행은 클라우드가 담당한다."),
        ("bullet", "공용 agent 시스템 구축: 회사나 팀이 하나의 표준 agent 구성을 운영하여 반복 설치와 설정 불일치를 줄인다."),
        ("bullet", "상시 가동성 확보: 클라우드 환경에서 OpenClaw Gateway와 API Gateway를 운영해 agent 호출 가능 시간을 늘린다."),
        ("bullet", "격리성과 운영성 향상: 사용자별 workspace와 container를 분리하고 실행 로그와 리소스 사용량을 추적한다."),
    ]
    add_cell_content(doc.tables[7].cell(0, 0), need)

    related_rows = [
        ("OpenClaw Gateway / OpenAI-compatible API", "OpenClaw는 Gateway에서 /v1/chat/completions, /v1/models 등 OpenAI-compatible HTTP surface를 제공하며 model: openclaw/<agentId> 형식으로 agent target을 라우팅할 수 있다. 다만 문서상 해당 엔드포인트는 operator 권한에 가까우므로 public SaaS용 per-user 권한 계층은 별도로 필요하다.", "OpenClaw docs/gateway/openai-http-api.md"),
        ("OpenClaw Docker Sandbox", "OpenClaw는 agent/session/shared scope의 sandbox를 제공하고 Docker backend에서 tool 실행과 workspace 접근을 컨테이너 안으로 제한할 수 있다. 기본 network none, read-only root, cap drop 등의 보안 설정을 기반으로 한다.", "OpenClaw docs/gateway/sandboxing.md"),
        ("OpenClaw Auth Profiles / Codex OAuth", "OpenClaw는 openai-codex OAuth login과 auth.order.openai 설정을 통해 OpenAI/Codex 모델 접근 profile을 선택할 수 있다. 본 프로젝트는 이를 사용자별 OAuth profile routing으로 확장한다.", "OpenClaw docs/providers/openai.md"),
        ("Docker Resource Constraints", "Docker는 컨테이너별 CPU, memory 등 리소스 제한을 제공한다. 이를 통해 agent sandbox별 자원 사용 상한과 사용량 추적을 설계할 수 있다.", "https://docs.docker.com/engine/containers/resource_constraints/"),
        ("Serverless Container Services", "AWS Fargate, Google Cloud Run 같은 서비스는 컨테이너 단위 실행과 서버 관리 부담 감소를 제공한다. 고도화 단계에서는 agent runtime을 VM Docker Compose에서 서버리스 컨테이너 실행으로 확장할 수 있다.", "AWS Fargate docs, Google Cloud Run docs"),
    ]
    related = [
        ("h1", "선행기술 및 차별점"),
        ("p", "본 프로젝트는 OpenClaw의 agent routing, sandbox, OAuth 기능을 기반으로 하지만, 기존 OpenClaw를 그대로 배포하는 것이 아니라 SaaS 형태의 멀티테넌트 서비스 계층을 추가한다. 특히 사용자별 workspace와 OAuth profile을 자동 선택하는 기능은 공용 agent 시스템을 실제 팀 서비스로 운영하기 위한 핵심 차별점이다."),
        ("table", (["기술/제품", "관련 기능", "본 프로젝트와의 관계"], related_rows, [2300, 4700, 2360])),
        ("h2", "기존 OpenClaw 대비 차별점"),
        ("bullet", "기존 OpenClaw는 agent별 workspace와 sandbox를 지원하지만, public API 사용자를 tenant로 구분하는 SaaS 권한 모델은 별도 구현이 필요하다."),
        ("bullet", "본 프로젝트는 공용 agent definition을 유지하면서 사용자별 workspace, session, OAuth profile, container usage를 분리한다."),
        ("bullet", "OpenClaw의 OpenAI-compatible API를 직접 외부에 공개하지 않고, 앞단에 인증/권한/사용량 추적을 담당하는 API Gateway를 둔다."),
    ]
    add_cell_content(doc.tables[9].cell(0, 0), related)

    architecture = [
        ("image", str(DIAGRAM)),
        ("h1", "구성 요소"),
        ("bullet", "사용자/API Client: 팀원이 브라우저, CLI, 또는 다른 백엔드에서 agent 실행을 요청한다."),
        ("bullet", "AaaS API Gateway: 사용자 인증, workspace 소유권 검증, rate limit, audit log, OpenClaw 호출 중계를 담당한다."),
        ("bullet", "Tenant/Workspace Router: user_id와 workspace_id를 기반으로 서버 내부의 /srv/agentaas/workspaces/{userId}/{workspaceId} 형태의 안전한 경로를 결정한다."),
        ("bullet", "OAuth Profile Router: user_id에 매핑된 openai-codex OAuth profile을 선택한다. 사용자는 profile id를 직접 지정하지 못하고, 서버가 검증된 profile만 OpenClaw 실행에 전달한다."),
        ("bullet", "OpenClaw Gateway: 공용 agent 정의, session routing, tool execution orchestration을 담당한다."),
        ("bullet", "Docker Sandbox Container: agent의 실행 도구와 workspace 접근이 이루어지는 격리 실행 환경이다. CPU/memory limit, network policy, idle cleanup을 적용한다."),
        ("h1", "동작 흐름"),
        ("num", "사용자가 API token과 workspace_id를 포함하여 agent 실행 요청을 보낸다."),
        ("num", "API Gateway는 token을 검증하고 해당 사용자가 workspace_id를 사용할 권한이 있는지 확인한다."),
        ("num", "Tenant Router는 사용자별 workspace 경로와 session key를 생성하고, OAuth Router는 사용자별 OpenAI/Codex profile을 선택한다."),
        ("num", "OpenClaw Gateway는 model: openclaw/<agentId> 형태의 agent target으로 공용 agent 정의를 실행한다."),
        ("num", "실제 파일 접근과 tool execution은 사용자 workspace가 mount된 Docker sandbox container 안에서 수행된다."),
        ("num", "결과는 API Gateway를 통해 사용자에게 반환되며, 실행 시간과 리소스 사용량은 audit/usage log에 기록된다."),
    ]
    add_cell_content(doc.tables[11].cell(0, 0), architecture)

    env = [
        ("h1", "동작 환경"),
        ("p", "기본 구현 대상은 퍼블릭 클라우드 VM 위의 Docker 기반 배포이다. 초기 단계에서는 AWS EC2, GCP Compute Engine, Azure VM 등 하나의 Linux VM에 Docker Engine과 OpenClaw Gateway, AaaS API Gateway, 사용자별 workspace volume을 구성한다. 이를 통해 로컬 PC 자원 부담을 줄이면서도 구현 난이도를 수업 프로젝트 범위 안에 유지한다."),
        ("h2", "기본 배포"),
        ("bullet", "Cloud VM: OpenClaw Gateway, API Gateway, Docker daemon, workspace volume, log storage를 운영한다."),
        ("bullet", "Docker Sandbox: agent 실행 시 사용자별 workspace를 mount하고 CPU/memory/network 정책을 적용한다."),
        ("bullet", "Reverse Proxy: HTTPS, API token 인증, 요청 크기 제한, rate limit을 담당한다."),
        ("bullet", "Local Client: 사용자는 브라우저, curl, Python client, 또는 OpenAI-compatible client를 통해 agent를 호출한다."),
        ("h2", "확장 배포"),
        ("bullet", "컨테이너 사용량 측정과 idle cleanup을 구현한 뒤, ECS Fargate, Cloud Run, Azure Container Apps 같은 서버리스 컨테이너 서비스로 agent sandbox 실행을 확장한다."),
        ("bullet", "이 경우 VM 단위 고정 비용이 아니라 container task 실행 시간과 할당 리소스 기준의 비용 모델에 더 가까워진다."),
        ("h2", "개발 범위"),
        ("bullet", "1차 목표: Cloud VM + Docker Compose 기반 POC, 두 명 이상의 사용자, 두 개 이상의 workspace, 사용자별 OAuth profile routing 검증"),
        ("bullet", "2차 목표: agent별/사용자별 사용량 대시보드, 컨테이너 리소스 제한 및 자동 정리, mock LLM 기반 테스트"),
        ("bullet", "제외 범위: 로컬 workspace bridge, 모바일/라즈베리파이 실행, 완전한 상용 billing 시스템"),
    ]
    add_cell_content(doc.tables[13].cell(0, 0), env)

    doc.save(str(PLAN_DOCX))


def build_handout() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    style = doc.styles["Normal"]
    style.font.name = FONT_LATIN
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.25

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("팀원용 개발 핸드아웃")
    apply_font(r, 22, BLACK, True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run("OpenClaw 기반 Multi-Tenant Agent-as-a-Service 플랫폼")
    apply_font(r2, 13, GRAY, False)

    add_callout(
        doc,
        "한 줄 목표",
        "공용 OpenClaw agent 시스템을 클라우드에 올리고, 각 사용자는 자기 workspace와 OAuth profile로 격리된 컨테이너 실행 환경에서 agent를 API처럼 호출한다.",
    )

    add_heading(doc, "개발 목표", 1)
    for item in [
        "로컬 PC에서 여러 agent/container를 실행할 때 발생하는 CPU/메모리 부담을 클라우드 실행으로 이전한다.",
        "팀/회사 사람들이 동일한 agent 정의를 공유하되, 작업 파일과 모델 인증은 사용자별로 분리한다.",
        "사용자별 workspace, OAuth profile, session, sandbox container를 묶어 안전한 실행 단위로 관리한다.",
        "컨테이너 실행 시간과 리소스 사용량을 기록하여 종량제 비용 분석의 기반을 만든다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "핵심 구성", 1)
    add_simple_table(
        doc,
        ["모듈", "역할"],
        [
            ("AaaS API Gateway", "사용자 인증, agent 호출 API, rate limit, audit log"),
            ("Tenant Workspace Router", "user_id + workspace_id를 안전한 서버 내부 경로로 매핑"),
            ("OAuth Profile Router", "사용자별 OpenAI/Codex OAuth profile 선택"),
            ("OpenClaw Gateway", "공용 agent 정의 실행, session/agent routing"),
            ("Docker Sandbox", "사용자 workspace를 mount하고 tool execution을 격리"),
            ("Usage Monitor", "요청 수, 실행 시간, CPU/memory, 사용자별 사용량 기록"),
        ],
        [2400, 6960],
    )

    add_heading(doc, "개발 Outline", 1)
    for item in [
        "1주차: OpenClaw fork 분석, cloud VM 개발 환경 구성, Docker sandbox 동작 확인",
        "2주차: API Gateway POC 구현, agent target 호출, 사용자 token 검증",
        "3주차: 사용자별 workspace router 구현 및 path traversal 방어",
        "4주차: OAuth profile router 구현, 사용자별 OpenAI/Codex profile 선택 검증",
        "5주차: container resource limit, idle cleanup, usage/audit log 구현",
        "6주차: mock LLM 테스트, 데모 시나리오, 최종 보고서/발표자료 정리",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "데모 시나리오", 1)
    for item in [
        "User A와 User B가 같은 code-review agent를 호출한다.",
        "각 사용자는 서로 다른 workspace에서 작업하고, 서로의 파일을 읽을 수 없다.",
        "각 agent 실행은 별도 Docker sandbox에서 수행되며 리소스 제한이 적용된다.",
        "모델 호출은 사용자별 OAuth profile로 수행되고, 사용량 로그가 사용자별로 남는다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "팀 내 작업 분담 초안", 1)
    add_simple_table(
        doc,
        ["역할", "담당 내용"],
        [
            ("플랫폼/배포", "Cloud VM, Docker Compose, reverse proxy, OpenClaw Gateway 실행"),
            ("API/인증", "API Gateway, token 인증, workspace 소유권 검증"),
            ("OpenClaw 연동", "agent target 호출, workspaceDir/authProfileId 주입 방식 조사 및 구현"),
            ("관측성/테스트", "usage log, mock LLM, 격리/권한 테스트, 데모 데이터"),
        ],
        [2400, 6960],
    )

    doc.save(str(HANDOUT_DOCX))


if __name__ == "__main__":
    build_plan_doc()
    build_handout()
    print(PLAN_DOCX)
    print(HANDOUT_DOCX)
